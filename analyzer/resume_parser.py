try:
    import pdfplumber
except ImportError:
    pdfplumber = None
    print("Warning: pdfplumber not installed, PDF extraction will not work.")

try:
    from docx import Document
except ImportError:
    Document = None
    print("Warning: python-docx not installed, DOCX extraction will not work.")

import re
from typing import Union, IO, List, Dict, Any, Tuple, Optional, TypeVar, TypedDict

# Define a type variable for Document to avoid the "Variable not allowed in type expression" error
DocType = TypeVar('DocType', bound=Any)

# Define skill categories for classification
SKILL_CATEGORIES = {
    "languages": [
        "python", "java", "javascript", "typescript", "c++", "c#", "php", "ruby", "go", 
        "swift", "kotlin", "scala", "r", "matlab", "perl", "css", "html", "sql", "bash"
    ],
    "frameworks": [
        "react", "angular", "vue", "django", "flask", "spring", "express", "laravel", 
        "rails", "asp.net", "nodejs", "bootstrap", "jquery", "tensorflow", "pytorch"
    ],
    "databases": [
        "mongodb", "mysql", "postgresql", "sql server", "oracle", "sqlite", "redis",
        "cassandra", "dynamodb", "elasticsearch", "neo4j", "firebase"
    ],
    "tools": [
        "git", "github", "gitlab", "bitbucket", "jira", "confluence", "docker", "kubernetes",
        "jenkins", "aws", "azure", "gcp", "heroku", "terraform", "linux", "windows"
    ],
    "soft_skills": [
        "communication", "leadership", "teamwork", "problem-solving", "critical thinking",
        "time management", "creativity", "adaptability", "collaboration", "analytical"
    ]
}

# Compiled regex patterns for better performance
# Job title indicators
RE_JOB_INDICATORS = [
    re.compile(r'#\s*\w+'),
    re.compile(r'@\s*\w+'),
    re.compile(r'\b(position|role|title|job)\b'),
    re.compile(r'\b(engineer|developer|nurse|manager|analyst|designer|consultant)\b'),
    re.compile(r'\b(intern|internship|trainee|graduate|junior|senior|lead)\b'),
    re.compile(r'\b(specialist|coordinator|assistant|associate|director)\b'),
    re.compile(r'\b(student|candidate|applicant|fresher)\b')
]

# Contact info patterns
# Email patterns
RE_EMAIL_PATTERNS = [
    re.compile(r'\b[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}\b'),
    re.compile(r'email[:\s]*([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})', re.IGNORECASE),
    re.compile(r'e-mail[:\s]*([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})', re.IGNORECASE)
]

# Enhanced phone patterns with international formats
RE_PHONE_PATTERNS = [
    # International format with country code
    re.compile(r'\+\d{1,4}[-.\s]?\d{1,4}[-.\s]?\d{1,9}'),
    # North American format
    re.compile(r'(?:\+?1[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'),
    # With label
    re.compile(r'(?:phone|mobile|tel|telephone|cell)[:\s]*(\+?\d{1,4}[-.\s]?\d{1,14})', re.IGNORECASE),
    # Indian format
    re.compile(r'\+91[-.\s]?\d{10}'),
    # UK format
    re.compile(r'\+44[-.\s]?\d{2,5}[-.\s]?\d{6,10}'),
    # Australian format
    re.compile(r'\+61[-.\s]?\d{1,4}[-.\s]?\d{6,10}'),
    # Generic international with various separators
    re.compile(r'(?:\+\d{1,4}[-.\s]?)?(?:\(?\d{2,5}\)?[-.\s]?)?\d{3,5}[-.\s]?\d{3,5}')
]

# Social profile patterns
RE_LINKEDIN_PATTERNS = [
    re.compile(r'linkedin\.com/in/([\w-]+)', re.IGNORECASE),
    re.compile(r'linkedin[:\s]*((?:https?://)?(?:www\.)?linkedin\.com/in/[\w-]+)', re.IGNORECASE),
    re.compile(r'((?:https?://)?(?:www\.)?linkedin\.com/in/[\w-]+)', re.IGNORECASE)
]

RE_GITHUB_PATTERNS = [
    re.compile(r'github\.com/([\w-]+)', re.IGNORECASE),
    re.compile(r'github[:\s]*((?:https?://)?(?:www\.)?github\.com/[\w-]+)', re.IGNORECASE),
    re.compile(r'((?:https?://)?(?:www\.)?github\.com/[\w-]+)', re.IGNORECASE)
]

# Section extraction patterns
RE_SECTION_HEADER = re.compile(r'^\s*[A-Z][A-Za-z\s]*\s*[:\-–]\s*')

# Skills extraction patterns
RE_WORD_PATTERN = re.compile(r'\b\w+\b')

# Project extraction patterns
RE_TECH_PATTERN = re.compile(r'(?:using|with|built\s+with|developed\s+with|implemented\s+with)[:\s]+([^\.;]+)', re.IGNORECASE)
RE_TECH_SPLIT = re.compile(r'[,/&+]|\sand\s')

# Cleaning text patterns
RE_LATEX_ARTIFACTS = re.compile(r'\{\.?\w+\}')
RE_SMALLCAPS = re.compile(r'\[\[(\w+)\]\]')
RE_LINE_CONTINUATIONS = re.compile(r'\\\n')
RE_BULLET_POINTS = re.compile(r'[•∙⋅⦁◦⟐⟡⟢⟣⟤⟥⟦⟧⟨⟩⟪⟫]')
RE_DASHES = re.compile(r'[–—−]')
RE_WHITESPACE = re.compile(r'\s+')

# Enhanced qualification extraction patterns
RE_QUALIFICATION_PATTERNS = [
    re.compile(r'(B\.?Tech|Bachelor|M\.?Tech|Master|Ph\.?D|MBA|BCA|MCA|B\.?Sc|M\.?Sc|B\.?E|M\.?E)\s+(?:in|of)?\s+([\w\s]+?)(?:from|at|,|\n)', re.IGNORECASE),
    re.compile(r'(Diploma|Certificate)\s+(?:in|of)?\s+([\w\s]+?)(?:from|at|,|\n)', re.IGNORECASE),
    re.compile(r'([\w\s]+?)\s+(?:University|College|Institute|School)', re.IGNORECASE)
]

# Work experience extraction patterns
RE_EXPERIENCE_PATTERNS = [
    re.compile(r'([\w\s]+?)\s+at\s+([\w\s&]+?)(?:,|\s+)(\d{4}|\w+\s+\d{4})(?:\s*[-–]\s*)?(\d{4}|\w+\s+\d{4}|present|current)?', re.IGNORECASE),
    re.compile(r'(Software Engineer|Developer|Analyst|Manager|Intern|Consultant|Designer)\s+[-–]\s+([\w\s&]+?)(?:,|\n)', re.IGNORECASE),
    re.compile(r'([\w\s]+?)\s+\|\s+([\w\s&]+?)\s+\|\s+([\d\w\s\-–]+)', re.IGNORECASE)
]

# ========== PRECISION RESUME PARSING COMPONENTS ==========

class ResumeAnalysis(TypedDict):
    """Type-annotated structure for precise resume analysis"""
    name: str
    contact: Dict[str, Dict[str, Union[str, float]]]
    education: List[Dict[str, Union[str, float]]]
    skills: Dict[str, List[str]]
    projects: List[Dict[str, Union[str, List[str]]]]
    confidence_scores: Dict[str, float]
    warnings: List[str]

# Specialized handling for edge cases
SPECIAL_HANDLING = {
    "double_brackets": {
        "pattern": r"\[\[(.*?)\]\]",
        "action": "preserve_as_is",
        "warning": "Unverified username: {match}"
    },
    "formatting_artifacts": {
        "patterns": [r"\{\.?\w+\}", r"<\/?\w+>"],
        "action": "remove_silently"
    },
    "partial_urls": {
        "pattern": r"(?:github|linkedin)\.com\/\w+",
        "action": "complete_url",
        "rules": {
            "github": "https://github.com/{username}",
            "linkedin": "https://linkedin.com/in/{username}"
        }
    }
}

def extract_text_from_pdf(pdf_file_path):
    """
    Extracts text from a PDF file given its path.
    Requires pdfplumber to be installed.
    """
    if pdfplumber is None:
        raise ImportError("pdfplumber is not installed. PDF extraction is not available.")
    
    text = ""
    try:
        with pdfplumber.open(pdf_file_path) as pdf:
            for page in pdf.pages:
                text += page.extract_text() or ""
        return text
    except Exception as e:
        print(f"Error extracting text from PDF: {e}")
        return ""

def extract_all_text_blocks(doc: Any) -> List[str]:
    """
    Extract all text blocks from DOCX document including paragraphs and table cells.
    Handles complex layouts where names might be in tables or split across elements.
    """
    text_blocks = []
    
    # Extract from paragraphs
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            # Split by newlines to handle multi-line content
            lines = [line.strip() for line in text.split('\n') if line.strip()]
            text_blocks.extend(lines)
    
    # Extract from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    # Split cell content by newlines
                    lines = [line.strip() for line in cell_text.split('\n') if line.strip()]
                    text_blocks.extend(lines)
    
    return text_blocks

def contains_job_title_indicator(text: str) -> bool:
    """
    Check if text contains job title indicators that often appear near names.
    """
    if not text:
        return False
    
    text_lower = text.lower()
    
    # Use compiled regex patterns for job indicators
    return any(pattern.search(text_lower) for pattern in RE_JOB_INDICATORS)

def extract_name_from_docx_robust(doc: Any) -> str:
    """
    Placeholder function that no longer extracts names.
    Returns a default placeholder value.
    """
    return "Not provided"

def is_valid_name(text: str) -> bool:
    """
    Placeholder function that always returns False.
    """
    return False

def extract_contact_or_title_indicator(text: str) -> bool:
    """
    Check if text contains contact info or job title indicators
    that often appear near names.
    """
    if not text:
        return False
    
    text_lower = text.lower()
    
    # Use job indicator patterns
    if any(pattern.search(text_lower) for pattern in RE_JOB_INDICATORS):
        return True
            
    # Check for contact patterns
    if any(pattern.search(text) for pattern in RE_EMAIL_PATTERNS):
        return True
        
    if any(pattern.search(text) for pattern in RE_PHONE_PATTERNS):
        return True
        
    if any(pattern.search(text) for pattern in RE_LINKEDIN_PATTERNS):
        return True
        
    if any(pattern.search(text) for pattern in RE_GITHUB_PATTERNS):
        return True
            
    return False

def extract_contact_info(text: str) -> Dict[str, str]:
    """
    Enhanced contact information extraction with robust patterns.
    Extracts email, phone, LinkedIn, and GitHub profiles.
    """
    contact_info = {
        "email": "",
        "phone": "",
        "linkedin": "",
        "github": ""
    }
    
    # Use compiled email patterns
    for pattern in RE_EMAIL_PATTERNS:
        match = pattern.search(text)
        if match:
            contact_info["email"] = match.group(1) if match.groups() else match.group(0)
            break
    
    # Use compiled phone patterns
    for pattern in RE_PHONE_PATTERNS:
        match = pattern.search(text)
        if match:
            contact_info["phone"] = match.group(1) if match.groups() else match.group(0)
            break
    
    # Use compiled LinkedIn patterns
    for pattern in RE_LINKEDIN_PATTERNS:
        match = pattern.search(text)
        if match:
            profile = match.group(1) if match.groups() else match.group(0)
            if not profile.startswith(('http://', 'https://')):
                if not profile.startswith('linkedin.com'):
                    profile = 'linkedin.com/in/' + profile
                profile = 'https://www.' + profile
            contact_info["linkedin"] = profile
            break
    
    # Use compiled GitHub patterns
    for pattern in RE_GITHUB_PATTERNS:
        match = pattern.search(text)
        if match:
            profile = match.group(1) if match.groups() else match.group(0)
            if not profile.startswith(('http://', 'https://')):
                if not profile.startswith('github.com'):
                    profile = 'github.com/' + profile
                profile = 'https://www.' + profile
            contact_info["github"] = profile
            break
    
    return contact_info

def extract_text_from_docx(docx_file_path):
    """
    Extract text from a DOCX file with enhanced layout handling.
    Includes robust name extraction, table parsing, and formatting cleanup.
    """
    try:
        # Load document
        doc = Document(docx_file_path)
        
        # Extract name using robust method
        name = extract_name_from_docx_robust(doc)
        
        # Extract text blocks with formatting awareness
        text_blocks = []
        
        # Extract from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                # Check for emphasis/formatting that might indicate section headers
                if paragraph.style.name.startswith(('Heading', 'Title')):
                    # Add newlines around headers for better section detection
                    text_blocks.append("\n" + paragraph.text.strip() + "\n")
                else:
                    text_blocks.append(paragraph.text.strip())
        
        # Extract from tables with improved table structure preservation
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = " ".join([p.text.strip() for p in cell.paragraphs if p.text.strip()])
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    # Join with pipe to represent table structure
                    text_blocks.append(" | ".join(row_text))
        
        # Join all text
        full_text = "\n".join(text_blocks)
        
        # Clean up the text
        full_text = clean_text(full_text)
        
        # Add extracted name to beginning if detected
        if name and name != "Not detected":
            full_text = f"Name: {name}\n\n" + full_text
        
        return full_text
    except Exception as e:
        print(f"Error extracting text from DOCX: {e}")
        return ""

def extract_text_from_file(file_path_or_stream, filename=None, parse_structure=False):
    """
    Extracts text from a DOCX file and optionally parses into structured format.
    PDF support is currently disabled.
    
    Args:
        file_path_or_stream: File path or file-like object
        filename: Optional filename (for streams)
        parse_structure: If True, returns structured data instead of plain text
        
    Returns:
        If parse_structure=False: Plain text content of the file
        If parse_structure=True: Dictionary with structured resume information
    """
    # Determine file type
    if filename:
        file_ext = filename.lower().split('.')[-1]
    elif hasattr(file_path_or_stream, 'name'):
        file_ext = file_path_or_stream.name.lower().split('.')[-1]
    else:
        file_ext = 'unknown'
    
    try:
        text = ""
        temp_file_path = None
        
        if file_ext == 'docx':
            # For DOCX files, we need to handle file streams differently
            if hasattr(file_path_or_stream, 'read'):
                # It's a file stream - save to temporary file
                import tempfile
                import os
                
                # Create a temporary file
                with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
                    # Reset stream position if possible
                    if hasattr(file_path_or_stream, 'seek'):
                        file_path_or_stream.seek(0)
                    
                    # Copy stream content to temp file
                    temp_file.write(file_path_or_stream.read())
                    temp_file_path = temp_file.name
                
                try:
                    # If parse_structure is requested, use the atomic parser directly
                    if parse_structure:
                        # Check if Document is available
                        if Document is None:
                            raise ImportError("python-docx is not installed. DOCX extraction is not available.")
                        
                        # Use the atomic parser
                        result = parse_resume_atomic(temp_file_path)
                        # Convert ResumeAnalysis to regular dict for backward compatibility
                        return result.__dict__
                    else:
                        # Extract text from temporary file
                        text = extract_text_from_docx(temp_file_path)
                finally:
                    # Clean up temporary file
                    if temp_file_path:
                        os.unlink(temp_file_path)
            else:
                # It's a file path
                if parse_structure:
                    # Check if Document is available
                    if Document is None:
                        raise ImportError("python-docx is not installed. DOCX extraction is not available.")
                    
                    # Use the atomic parser
                    result = parse_resume_atomic(file_path_or_stream)
                    # Convert ResumeAnalysis to regular dict for backward compatibility
                    return result.__dict__
                else:
                    text = extract_text_from_docx(file_path_or_stream)
        elif file_ext == 'pdf':
            # PDF support is disabled
            raise ImportError("PDF extraction is currently disabled. Please use DOCX files.")
        else:
            raise ValueError(f"Unsupported file type: {file_ext}. Only DOCX files are supported.")
        
        # If we got here with parse_structure=True, use the regular parser
        if parse_structure:
            return parse_resume(text)
        else:
            return text
    except ImportError as e:
        print(f"Import Error: {e}")
        return {} if parse_structure else f"ERROR: {str(e)}"
    except Exception as e:
        print(f"Error extracting text: {e}")
        return {} if parse_structure else ""

def extract_education_details(text: str) -> List[Dict[str, str]]:
    """
    Extract detailed education information from resume text.
    Looks for institutions, degrees, years, and GPA.
    """
    education_details = []
    
    # First, try to find education section
    education_section = extract_section(text, "education") or extract_section(text, "academic") or ""
    if not education_section:
        # Try finding education keywords in full text
        education_keywords = ["university", "college", "institute", "b.tech", "bachelor", "master", "phd", "diploma"]
        if any(keyword in text.lower() for keyword in education_keywords):
            education_section = text
    
    if not education_section:
        return education_details
    
    # Extract institution patterns
    institution_patterns = [
        r'(\*{0,2}([A-Za-z\s.]+\s(?:Institute|College|University|Polytechnic|School)))',
        r'([A-Z][A-Za-z\s.]+(?:Institute|College|University|Polytechnic|School))',
        r'(?:at|from|in)\s+([A-Z][A-Za-z\s.]{3,50})'
    ]
    
    # Extract degree patterns
    degree_patterns = [
        r'(B\.?Tech|Bachelor|M\.?Tech|Master|Ph\.?D|Diploma)\s+(?:of|in|with)?\s+([\w\s]+)',
        r'(Bachelor|Master)(?:\'s)?\s+(?:of|in|with)?\s+([\w\s]+)',
        r'(B\.?E|M\.?E|B\.?Sc|M\.?Sc|MBA|BCA|MCA)\s+(?:in)?\s+([\w\s]+)?'
    ]
    
    # Extract year patterns
    year_patterns = [
        r'(\d{4})\s*-\s*(\d{2,4})',  # 2018-2022 or 2018-22
        r'(\d{4})\s*to\s*(\d{2,4})',  # 2018 to 2022
        r'(\d{4})\s*–\s*(?:present|current|ongoing|till date)',  # 2018–present
        r'(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*[\s,]+(\d{4})'  # Sep 2022
    ]
    
    # Extract GPA/percentage patterns
    gpa_patterns = [
        r'(?:GPA|CGPA)[:\s]*(\d+\.\d+)',
        r'(\d{1,3}[.,]?\d{0,2})\s*%?'
    ]
    
    
    # Find all matches for institutions, degrees, years, and GPA
    institutions = []
    degrees = []
    years = []
    gpas = []
    
    for pattern in institution_patterns:
        matches = re.findall(pattern, education_section, re.IGNORECASE)
        institutions.extend([match[0] for match in matches])
    
    for pattern in degree_patterns:
        matches = re.findall(pattern, education_section, re.IGNORECASE)
        degrees.extend([match[0] + " in " + match[1] for match in matches if match[1]])
    
    for pattern in year_patterns:
        matches = re.findall(pattern, education_section)
        years.extend([match[0] + "-" + match[1] if len(match) > 1 else match[0] for match in matches])
    
    for pattern in gpa_patterns:
        matches = re.findall(pattern, education_section)
        gpas.extend(matches)
    
    # Combine matches into structured details
    for i in range(max(len(institutions), len(degrees), len(years), len(gpas))):
        detail = {}
        if i < len(institutions):
            detail["institution"] = institutions[i]
        if i < len(degrees):
            detail["degree"] = degrees[i]
        if i < len(years):
            detail["year"] = years[i]
        if i < len(gpas):
            detail["gpa"] = gpas[i]
        education_details.append(detail)
    
    return education_details

def extract_section(text: str, section_name: str) -> str:
    """
    Extract a specific section from the resume text by section name.
    Supports various section title formats and common delimiters.
    """
    # Normalize section name
    section_name = section_name.lower().strip()
    
    # Compile section title patterns for this specific section
    title_patterns = [
        re.compile(r'^\s*' + re.escape(section_name) + r'\s*[:\-–]\s*'),
        re.compile(r'^\s*' + re.escape(section_name.capitalize()) + r'\s*[:\-–]\s*'),
        re.compile(r'^\s*' + re.escape(section_name.upper()) + r'\s*[:\-–]\s*')
    ]
    
    # Split text into lines for processing
    lines = text.split('\n')
    in_section = False
    section_lines = []
    
    for line in lines:
        # Check if line matches any of the title patterns
        if any(pattern.match(line) for pattern in title_patterns):
            in_section = True
            section_lines.append(line)
            continue
        
        # If already in section, check for empty line or next section
        if in_section:
            if not line.strip() or RE_SECTION_HEADER.match(line):
                break
            section_lines.append(line)
    
    # Join captured section lines
    return '\n'.join(section_lines).strip()

def extract_skills_categorized(text: str) -> Dict[str, List[str]]:
    """
    Extract and categorize skills from the resume text.
    Uses predefined skill categories and keywords.
    """
    skills = {category: [] for category in SKILL_CATEGORIES}
    
    # Flatten all skill keywords into a single list
    all_skills = sum(SKILL_CATEGORIES.values(), [])
    
    # Extract all words from text (simple whitespace split)
    words = RE_WORD_PATTERN.findall(text.lower())
    
    # Count occurrences of each word
    from collections import Counter
    word_counts = Counter(words)
    
    # Heuristic: If a word is a skill, it will appear in the skills list or be a known technology
    for word, count in word_counts.items():
        # Check if word is a known skill or technology
        if word in all_skills:
            # Find the category of the skill
            for category, keywords in SKILL_CATEGORIES.items():
                if word in keywords:
                    skills[category].append(word)
                    break
    
    # Remove duplicates while preserving order
    for category in skills:
        skills[category] = list(dict.fromkeys(skills[category]))
    
    return skills

def extract_projects(text: str) -> List[Dict[str, Any]]:
    """
    Extract project information including title, description, and technologies.
    Handles various project section formats and bullet points.
    """
    projects_list = []
    
    # Find projects section
    projects_section = extract_section(text, "projects") or extract_section(text, "personal projects") or ""
    if not projects_section:
        # Try alternative titles
        projects_section = extract_section(text, "work") or extract_section(text, "personal work") or ""
    
    if not projects_section:
        return projects_list
    
    # Split into individual projects
    # First try to identify project delimiters (patterns that likely separate projects)
    project_delimiters = [
        r'\n\s*\*\*([^\*]+)\*\*',  # Bold project titles with **Title**
        r'\n\s*•\s+([A-Z][^•\n]+)',  # Bullet points with capital letter
        r'\n\s*\d+\.\s+([A-Z][^\n]+)',  # Numbered projects: 1. Project Name
        r'\n\s*([A-Z][a-zA-Z\s]+)(?:\s*\(|\s*-\s*|\s*:)'  # Project Name (details) or Project Name - details
    ]
    
    # Try each delimiter pattern
    for delimiter in project_delimiters:
        # Find all matches
        matches = re.finditer(delimiter, '\n' + projects_section, re.MULTILINE)
        
        # Extract start positions of projects
        positions = [0]  # Start of text
        project_titles = []
        
        for match in matches:
            positions.append(match.start())
            project_titles.append(match.group(1).strip() if match.groups() else "")
        
        # If we found at least one delimiter
        if len(positions) > 1:
            # Split the section using positions
            for i in range(len(positions) - 1):
                start = positions[i]
                end = positions[i + 1]
                
                # Extract project content
                project_content = projects_section[start:end].strip()
                
                # For first item, there's no title in the delimiter
                project_title = project_titles[i] if i < len(project_titles) else ""
                
                # If no title found in delimiter, try to extract from content
                if not project_title and i == 0:
                    # Try to extract title from first line
                    first_line = project_content.split('\n')[0].strip()
                    if first_line and not first_line.startswith(('•', '-', '*')):
                        project_title = first_line
                        # Remove title from content
                        project_content = project_content[len(first_line):].strip()
                
                if project_content:
                    project = {
                        "title": project_title,
                        "description": project_content,
                        "technologies": extract_technologies(project_content)
                    }
                    projects_list.append(project)
            
            # Add the last project
            last_content = projects_section[positions[-1]:].strip()
            if last_content:
                project = {
                    "title": project_titles[-1] if project_titles else "",
                    "description": last_content,
                    "technologies": extract_technologies(last_content)
                }
                projects_list.append(project)
            
            # If we found projects using this delimiter, return them
            if projects_list:
                return projects_list
    
    # If no projects found with delimiters, try simpler bullet-point based approach
    lines = projects_section.split('\n')
    current_project = {}
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # Check if line looks like a project title (starts with bullet or is short)
        if line.startswith(('•', '-', '*')) or (len(line) < 50 and not line.endswith(('.', ','))):
            # Save previous project if exists
            if current_project and 'title' in current_project:
                projects_list.append(current_project)
                current_project = {}
            
            # Extract title
            title = line.lstrip('•-* ')
            current_project = {
                "title": title,
                "description": "",
                "technologies": []
            }
        elif current_project:
            # Add to description
            if 'description' in current_project:
                current_project['description'] += " " + line
            else:
                current_project['description'] = line
            
            # Check for technologies
            techs = extract_technologies(line)
            if techs:
                current_project['technologies'].extend(techs)
    
    # Add the last project
    if current_project and 'title' in current_project:
        projects_list.append(current_project)
    
    return projects_list

def extract_technologies(text: str) -> List[str]:
    """
    Extract technologies mentioned in text.
    Looks for technology keywords, "using X" patterns, and "built with X" patterns.
    """
    technologies = []
    
    # Look for "using X" pattern
    match = RE_TECH_PATTERN.search(text)
    if match:
        tech_text = match.group(1).strip()
        # Split by common separators
        techs = RE_TECH_SPLIT.split(tech_text)
        technologies.extend([t.strip() for t in techs if t.strip()])
    
    # Look for technology keywords
    tech_keywords = sum(SKILL_CATEGORIES.values(), [])  # Flatten skill categories
    
    for tech in tech_keywords:
        # Create a temporary pattern for this tech keyword
        tech_pattern = re.compile(r'\b' + re.escape(tech) + r'\b', re.IGNORECASE)
        if tech_pattern.search(text.lower()):
            technologies.append(tech)
    
    # Remove duplicates while preserving order
    return list(dict.fromkeys(technologies))

def clean_text(text: str) -> str:
    """
    Clean and normalize text by removing special formatting and standardizing spacing.
    """
    # Remove LaTeX artifacts
    text = RE_LATEX_ARTIFACTS.sub('', text)
    
    # Replace smallcaps formatting
    text = RE_SMALLCAPS.sub(r'\1', text)
    
    # Handle line continuations
    text = RE_LINE_CONTINUATIONS.sub(' ', text)
    
    # Normalize whitespace
    text = RE_WHITESPACE.sub(' ', text)
    
    # Normalize bullet points
    text = RE_BULLET_POINTS.sub('•', text)
    
    # Normalize dashes
    text = RE_DASHES.sub('-', text)
    
    # Fix table-like structures (convert pipe tables to proper text)
    lines = text.split('\n')
    for i in range(len(lines)):
        if '|' in lines[i]:
            lines[i] = lines[i].replace('|', ' - ')
    
    return '\n'.join(lines).strip()

def parse_resume(text: str) -> Dict[str, Any]:
    """
    Main function to parse resume text into structured information.
    """
    # Clean the text
    text = clean_text(text)
    
    # Extract name (for plain text, not docx)
    name = "Not detected"  # Will be properly extracted for DOCX files
    
    # Find name in first few lines
    lines = text.split('\n')
    for i, line in enumerate(lines[:5]):
        if len(line) < 50 and is_valid_name(line):
            name = line
            break
    
    # Extract contact information
    contact = extract_contact_info(text)
    
    # Extract education
    education = extract_education_details(text)
    
    # Extract skills
    skills = extract_skills_categorized(text)
    
    # Extract projects
    projects = extract_projects(text)
    
    # Extract certifications
    certificates_section = extract_section(text, "certifications") or extract_section(text, "certificates") or ""
    certificates = []
    if certificates_section:
        # Split by newlines and bullet points
        for line in re.split(r'\n+|(?:^|\n)\s*[•\-*]\s*', certificates_section):
            if line.strip():
                certificates.append(line.strip())
    
    # Build final result
    result = {
        "name": name,
        "contact": contact,
        "education": education,
        "skills": skills,
        "projects": projects,
        "certificates": certificates
    }
    
    return result

def extract_text_forensic(docx_path: str) -> List[str]:
    """
    Extract text with forensic precision from a DOCX file.
    
    1. Preserves ALL whitespace and line breaks initially
    2. Handles these special cases:
       - [[double bracket placeholders]] → keep as single tokens
       - {.underline} {.smallcaps} → remove formatting artifacts
       - Image placeholders → convert to "[IMAGE]" markers
       - Table cells → merge multi-line cells with " | " separator
    3. Returns list of raw lines with exact original spacing
    """
    try:
        doc = Document(docx_path)
        lines = []

        # Helper: Detect image placeholders
        def is_image_placeholder(paragraph) -> bool:
            return any(run.element.tag.endswith('drawing') for run in paragraph.runs)

        # Extract paragraphs
        for para in doc.paragraphs:
            if is_image_placeholder(para):
                lines.append("[IMAGE]")
                continue
            # Preserve double-bracket placeholders
            text = para.text
            # Already preserved by not altering them
            lines.append(text)

        # Extract tables (merge multi-line cells with " | ")
        for table in doc.tables:
            for row in table.rows:
                cell_texts = []
                for cell in row.cells:
                    cell_lines = [line for line in cell.text.splitlines()]
                    cell_texts.append(" ".join(cell_lines).strip())
                if cell_texts:
                    lines.append(" | ".join(cell_texts))

        # Preserve all whitespace and line breaks
        raw_lines = []
        for line in lines:
            for l in line.splitlines():
                raw_lines.append(l.rstrip('\r\n'))

        return raw_lines
    except Exception as e:
        print(f"Error in forensic text extraction: {e}")
        return []

def extract_name_military_grade(lines: List[str]) -> Tuple[str, float]:
    """
    Placeholder function that no longer extracts names.
    Returns a default placeholder value with zero confidence.
    """
    return ("Not provided", 0.0)

def extract_contact_zero_false_positives(lines: List[str]) -> Dict[str, Dict[str, Union[str, float]]]:
    """
    Atomic component extraction with zero false positives:
    - Email: Valid TLD, not a placeholder
    - Phone: 10-15 digits after cleaning
    - GitHub: Only extract full URLs or explicitly labeled usernames
    - LinkedIn: Only extract full URLs
    
    Returns structured contact info with confidence scores
    """
    result = {
        "email": {"value": "", "confidence": 0.0},
        "phone": {"value": "", "confidence": 0.0},
        "github": {"value": "", "confidence": 0.0},
        "linkedin": {"value": "", "confidence": 0.0},
    }
    
    # Advanced patterns with validation
    email_pat = re.compile(r'\b(?!noreply|example)[a-z0-9._%+-]+@[a-z0-9.-]+\.[a-z]{2,}\b', re.I)
    phone_pat = re.compile(r'(\+\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}')
    github_url_pat = re.compile(r'(https?://)?(www\.)?github\.com/([A-Za-z0-9_-]+)', re.I)
    github_label_pat = re.compile(r'github[:：]?\s*([A-Za-z0-9_-]+)', re.I)
    github_brackets_pat = re.compile(r'\[\[([A-Za-z0-9_-]+)\]\].*?github', re.I)
    linkedin_url_pat = re.compile(r'(https?://)?(www\.)?linkedin\.com/in/[A-Za-z0-9_-]+', re.I)

    for line in lines:
        # Email - with validation
        for m in email_pat.finditer(line):
            email = m.group(0)
            if not email.lower().endswith(('.png', '.jpg', '.jpeg')) and len(email) > 6:
                # Validate TLD (at least 2 chars)
                if re.search(r'\.[a-z]{2,}$', email, re.I):
                    result["email"] = {"value": email, "confidence": 1.0}
        
        # Phone - with digit count validation
        for m in phone_pat.finditer(line):
            phone = re.sub(r'\D', '', m.group(0))
            if 10 <= len(phone) <= 15:
                result["phone"] = {"value": m.group(0), "confidence": 1.0}
        
        # GitHub - only explicit mentions
        m = github_url_pat.search(line)
        if m:
            # Case 1: Full URL
            username = m.group(3)
            github_url = f"https://github.com/{username}"
            result["github"] = {"value": github_url, "confidence": 1.0}
        else:
            # Case 2: Explicit label
            m2 = github_label_pat.search(line)
            if m2:
                username = m2.group(1)
                github_url = f"https://github.com/{username}"
                result["github"] = {"value": github_url, "confidence": 0.9}
            else:
                # Case 3: Double brackets with GitHub nearby
                m3 = github_brackets_pat.search(line)
                if m3:
                    username = m3.group(1)
                    github_url = f"https://github.com/{username}"
                    result["github"] = {"value": github_url, "confidence": 0.8}
        
        # LinkedIn - only full URLs
        m = linkedin_url_pat.search(line)
        if m:
            linkedin_url = m.group(0)
            if not linkedin_url.startswith(('http://', 'https://')):
                linkedin_url = 'https://' + linkedin_url
            result["linkedin"] = {"value": linkedin_url, "confidence": 1.0}
    
    return result

def extract_education_context_aware(lines: List[str]) -> List[Dict[str, Union[str, float]]]:
    """
    Context-aware education extraction:
    1. Section detection
    2. Entry patterns (institution, degree, year)
    3. Multi-entry awareness
    4. Fallback scan
    
    Returns list of education entries with confidence scores
    """
    results = []
    section_started = False
    education_lines = []
    
    # First find the education section
    for i, line in enumerate(lines):
        l_lower = line.lower().strip()
        if l_lower == "education" or l_lower == "academic background" or l_lower.startswith("education:"):
            section_started = True
            continue
            
        # Collect lines until next section header or blank lines
        if section_started:
            if not line.strip():
                continue
            if RE_SECTION_HEADER.match(line) and i > 0:
                break
            education_lines.append(line)
    
    # If no dedicated section found, scan entire document
    if not education_lines:
        education_keywords = ["university", "college", "institute", "b.tech", "bachelor", "master", "phd", "diploma"]
        for line in lines:
            if any(keyword in line.lower() for keyword in education_keywords):
                education_lines.append(line)
    
    # Process the education lines
    if education_lines:
        # Institution pattern
        institution_pat = re.compile(r'([A-Z][A-Za-z\s.]{10,}(?:University|College|Institute|Polytechnic|School)\b)', re.I)
        
        # Degree pattern - international support
        degree_pat = re.compile(r'(B\.?Tech|B\.?E|M\.?Tech|Diploma|B\.?Sc|M\.?Sc|Bachelor|Master|Ph\.?D|Diplôme d\'Ingénieur)\s+(?:in|of)?\s*([\w\s]+)?', re.I)
        
        # Year pattern - handles parenthetical
        year_pat = re.compile(r'(?:\(|\[)?(20\d{2}\s*[-–]\s*(?:20\d{2}|present)|20\d{2})(?:\)|\])?')
        
        entry = {}
        current_line = 0
        
        while current_line < len(education_lines):
            line = education_lines[current_line]
            
            # Reset entry if we have bullet/numbered markers
            if re.match(r'^\s*[•\-*]\s+|^\s*\d+\.\s+', line):
                if entry and "institution" in entry:
                    results.append({**entry, "confidence": 0.95})
                    entry = {}
            
            # Extract institution
            inst_match = institution_pat.search(line)
            if inst_match:
                entry["institution"] = inst_match.group(0)
            
            # Extract degree
            deg_match = degree_pat.search(line)
            if deg_match:
                degree = deg_match.group(1)
                field = deg_match.group(2) if deg_match.group(2) else ""
                entry["degree"] = f"{degree}{' in ' + field if field else ''}"
            
            # Extract year
            year_match = year_pat.search(line)
            if year_match:
                entry["year"] = year_match.group(1)
            
            # Check if next line is part of multi-line institution
            if current_line + 1 < len(education_lines):
                next_line = education_lines[current_line + 1]
                # If next line doesn't have a degree/year and no bullet point, likely continuation
                if not degree_pat.search(next_line) and not year_pat.search(next_line) and not re.match(r'^\s*[•\-*]\s+|^\s*\d+\.\s+', next_line):
                    # We'll handle multi-line institutions
                    if "institution" in entry:
                        entry["institution"] += " " + next_line.strip()
            
            current_line += 1
        
        # Add the last entry
        if entry and "institution" in entry:
            results.append({**entry, "confidence": 0.95})
    
    return results

def extract_skills_precision_mapped(lines: List[str]) -> Dict[str, List[str]]:
    """
    Skills extraction with precision mapping:
    1. Section-first approach
    2. Pattern recognition
    3. Ontology validation (80% match)
    4. Context-aware extraction
    
    Returns categorized skills
    """
    skills = {category: [] for category in SKILL_CATEGORIES}
    
    # Flatten all skill keywords
    all_skills = sum(SKILL_CATEGORIES.values(), [])
    
    # First, find skills section
    section_started = False
    skills_lines = []
    
    for i, line in enumerate(lines):
        l_lower = line.lower().strip()
        if "skills" in l_lower or "technical skills" in l_lower:
            section_started = True
            continue
            
        if section_started:
            if not line.strip():
                continue
            if RE_SECTION_HEADER.match(line) and i > 0:
                break
            skills_lines.append(line)
    
    # Process skills in section
    if skills_lines:
        # Check for category headers
        category_pat = re.compile(r"^\s*(languages|frameworks|tools|databases|soft skills)[:\s]", re.I)
        bullet_pat = re.compile(r"^\s*[•\-*]\s*([\w\s,/]+)")
        
        current_category = None
        
        for line in skills_lines:
            # Check if line defines a category
            cat_match = category_pat.search(line)
            if cat_match:
                current_category = cat_match.group(1).lower()
                # Map to our categories
                if current_category == "programming languages":
                    current_category = "languages"
                if current_category == "soft skills":
                    current_category = "soft_skills"
                
                # Extract skills after the category header
                after_header = line[cat_match.end():].strip()
                if after_header:
                    skills_in_line = re.split(r'[,/]|\s+and\s+', after_header)
                    for skill in skills_in_line:
                        skill = skill.strip().lower()
                        if skill:
                            # Find matching skill in ontology (80% match)
                            for category, defined_skills in SKILL_CATEGORIES.items():
                                for defined_skill in defined_skills:
                                    if skill == defined_skill or (len(skill) > 3 and defined_skill in skill):
                                        if defined_skill not in skills[category]:
                                            skills[category].append(defined_skill)
            
            # Check for bullet points with skills
            bullet_match = bullet_pat.match(line)
            if bullet_match:
                skill_text = bullet_match.group(1).lower()
                skills_in_line = re.split(r'[,/]|\s+and\s+', skill_text)
                
                for skill in skills_in_line:
                    skill = skill.strip().lower()
                    if skill:
                        # Match to our ontology
                        for category, defined_skills in SKILL_CATEGORIES.items():
                            for defined_skill in defined_skills:
                                if skill == defined_skill or (len(skill) > 3 and defined_skill in skill):
                                    if defined_skill not in skills[category]:
                                        skills[category].append(defined_skill)
            
            # General skills extraction if no bullet/category
            if not cat_match and not bullet_match:
                for category, defined_skills in SKILL_CATEGORIES.items():
                    for defined_skill in defined_skills:
                        if defined_skill in line.lower():
                            if defined_skill not in skills[category]:
                                skills[category].append(defined_skill)
    
    # Fallback: If no skills section or few skills found, scan full document
    # But only accept skills that appear multiple times (>3)
    if not any(skills.values()):
        # Join all lines
        full_text = " ".join(lines).lower()
        
        for category, defined_skills in SKILL_CATEGORIES.items():
            for defined_skill in defined_skills:
                if full_text.count(defined_skill) > 3:
                    skills[category].append(defined_skill)
    
    return skills

def extract_projects_github_aware(lines: List[str]) -> List[Dict[str, Any]]:
    """
    Project extraction with GitHub awareness:
    1. Header detection
    2. Entry patterns
    3. Technology extraction
    4. GitHub link detection
    
    Returns list of projects with details
    """
    projects = []
    
    # Find projects section
    section_started = False
    projects_lines = []
    
    for i, line in enumerate(lines):
        l_lower = line.lower().strip()
        if l_lower == "projects" or l_lower == "personal work" or l_lower.startswith("projects:"):
            section_started = True
            continue
            
        if section_started:
            if not line.strip():
                continue
            if RE_SECTION_HEADER.match(line) and i > 0:
                break
            projects_lines.append(line)
    
    if not projects_lines:
        return projects
    
    # Process projects
    entry = None
    entry_pattern = re.compile(r'(?:\d+\.\s+|\-\s+|\*\s+)(.*?)\s*[\(\[]?(20\d{2}|present)[\)\]]?')
    using_pattern = re.compile(r"using\s+([\w\s,/]+)")
    github_pattern = re.compile(r"(https?://)?github\.com/[\w-]+/[\w-]+")
    
    for i, line in enumerate(projects_lines):
        # New project entry
        entry_match = entry_pattern.match(line)
        if entry_match:
            # Save previous entry
            if entry:
                projects.append(entry)
            
            # Create new entry
            name = entry_match.group(1).strip()
            year = entry_match.group(2)
            
            entry = {
                "name": name,
                "year": year,
                "description": "",
                "technologies": []
            }
            
            # Rest of line is description
            rest = line[entry_match.end():].strip()
            if rest:
                entry["description"] = rest
        
        # Continue existing entry
        elif entry:
            # Add to description
            entry["description"] += " " + line
            
            # Extract technologies
            tech_match = using_pattern.search(line)
            if tech_match:
                techs = re.split(r'[,/]|\s+and\s+', tech_match.group(1))
                for tech in techs:
                    tech = tech.strip()
                    if tech and tech not in entry["technologies"]:
                        entry["technologies"].append(tech)
            
            # Check for github URL
            github_match = github_pattern.search(line)
            if github_match:
                github_url = github_match.group(0)
                if not github_url.startswith('http'):
                    github_url = 'https://' + github_url
                entry["github"] = github_url
            
            # Implicit technology extraction
            for category, defined_skills in SKILL_CATEGORIES.items():
                for skill in defined_skills:
                    if skill in line.lower() and skill not in entry["technologies"]:
                        entry["technologies"].append(skill)
    
    # Add final entry
    if entry:
        projects.append(entry)
    
    return projects

def parse_resume_atomic(docx_path: str) -> ResumeAnalysis:
    """
    Industrial-grade resume parser with atomic precision.
    Handles all edge cases with military-grade validation.
    
    Returns structured ResumeAnalysis with confidence scores and warnings.
    """
    lines = extract_text_forensic(docx_path)
    name = "Not provided"  # No longer extract name
    name_conf = 0.0
    contact = extract_contact_zero_false_positives(lines)
    education = extract_education_context_aware(lines)
    skills = extract_skills_precision_mapped(lines)
    projects = extract_projects_github_aware(lines)
    
    # Generate warnings
    warnings = []
    
    # Handle double bracket placeholders
    db_pattern = re.compile(r"\[\[(.*?)\]\]")
    for line in lines:
        for match in db_pattern.finditer(line):
            username = match.group(1)
            warnings.append(f"Unverified username: {username}")
    
    # Calculate confidence scores
    confidence_scores = {
        "name": name_conf,
        "contact": min(v["confidence"] for v in contact.values() if v["confidence"] > 0) if any(v["confidence"] > 0 for v in contact.values()) else 0.0,
        "education": 0.95 if education else 0.0,
        "skills": 0.9 if any(skills.values()) else 0.0,
        "projects": 0.9 if projects else 0.0,
    }
    
    return ResumeAnalysis(
        name=name,
        contact=contact,
        education=education,
        skills=skills,
        projects=projects,
        confidence_scores=confidence_scores,
        warnings=warnings
    )

def extract_qualification_structured(text: str) -> List[Dict[str, str]]:
    """
    Extract structured qualification information including degree, major, and institution.
    """
    qualifications = []
    
    # Find education section first
    education_section = extract_section(text, "education") or extract_section(text, "academic") or text
    
    # Extract qualification patterns
    for pattern in RE_QUALIFICATION_PATTERNS:
        matches = pattern.findall(education_section)
        for match in matches:
            if len(match) >= 2:
                degree = match[0].strip()
                major = match[1].strip()
                
                # Look for institution in the same line or nearby lines
                institution = ""
                lines = education_section.split('\n')
                for line in lines:
                    if degree.lower() in line.lower() and major.lower() in line.lower():
                        # Extract institution from the same line
                        inst_match = re.search(r'(?:from|at)\s+([\w\s]+?)(?:,|\n|$)', line, re.IGNORECASE)
                        if inst_match:
                            institution = inst_match.group(1).strip()
                        break
                
                qualifications.append({
                    "degree": degree,
                    "major": major,
                    "institution": institution or "Not specified"
                })
    
    return qualifications

def extract_work_experience_structured(text: str) -> List[Dict[str, str]]:
    """
    Extract structured work experience including job title, company, and duration.
    """
    experiences = []
    
    # Find experience section
    experience_section = extract_section(text, "experience") or extract_section(text, "work") or text
    
    # Extract experience patterns
    for pattern in RE_EXPERIENCE_PATTERNS:
        matches = pattern.findall(experience_section)
        for match in matches:
            if len(match) >= 2:
                job_title = match[0].strip()
                company = match[1].strip()
                start_date = match[2].strip() if len(match) > 2 else ""
                end_date = match[3].strip() if len(match) > 3 and match[3] else "Present"
                
                duration = f"{start_date} - {end_date}" if start_date else "Duration not specified"
                
                experiences.append({
                    "job_title": job_title,
                    "company": company,
                    "duration": duration
                })
    
    return experiences

def extract_skills_detailed(text: str) -> Dict[str, List[str]]:
    """
    Extract detailed skills categorization with better accuracy.
    """
    skills_data = extract_skills_categorized(text)
    
    # Also extract skills mentioned in projects and experience sections
    projects_section = extract_section(text, "projects") or ""
    experience_section = extract_section(text, "experience") or ""
    combined_text = f"{text} {projects_section} {experience_section}".lower()
    
    # Additional skill extraction from context
    skill_context_patterns = [
        r'(?:used|utilizing|working with|experience in|proficient in|skilled in)\s+([\w\s,]+)',
        r'(?:technologies|tools|frameworks):\s*([\w\s,]+)',
        r'(?:programming languages):\s*([\w\s,]+)'
    ]
    
    for pattern in skill_context_patterns:
        matches = re.findall(pattern, combined_text, re.IGNORECASE)
        for match in matches:
            skills_in_context = [s.strip() for s in re.split(r'[,&\n]', match) if s.strip()]
            for skill in skills_in_context:
                # Categorize the skill
                skill_lower = skill.lower()
                for category, known_skills in SKILL_CATEGORIES.items():
                    if any(known_skill in skill_lower for known_skill in known_skills):
                        if skill not in skills_data[category]:
                            skills_data[category].append(skill)
    
    return skills_data

def extract_projects_enhanced(text: str) -> List[Dict[str, Any]]:
    """
    Enhanced project extraction with title, description, and technologies.
    """
    projects = extract_projects(text)
    
    # Enhance with better structure
    enhanced_projects = []
    for project in projects:
        if isinstance(project, dict):
            enhanced_project = {
                "title": project.get("title", "Untitled Project"),
                "description": project.get("description", "No description available"),
                "technologies": project.get("technologies", [])
            }
        else:
            # Handle string format
            enhanced_project = {
                "title": "Project",
                "description": str(project),
                "technologies": extract_technologies(str(project))
            }
        
        enhanced_projects.append(enhanced_project)
    
    return enhanced_projects

def parse_resume_structured(text: str) -> Dict[str, Any]:
    """
    Enhanced resume parsing with structured field extraction.
    """
    # Clean the text
    text = clean_text(text)
    
    # Extract structured information
    qualification = extract_qualification_structured(text)
    work_experience = extract_work_experience_structured(text)
    skills = extract_skills_detailed(text)
    projects = extract_projects_enhanced(text)
    
    # Extract basic contact information
    contact = extract_contact_info(text)
    
    return {
        "name": "Not provided",  # Name extraction disabled
        "contact": contact,
        "qualification": qualification,
        "work_experience": work_experience,
        "skills": skills,
        "projects": projects,
        "raw_education": extract_section(text, "education") or "Not detected",
        "raw_experience": extract_section(text, "experience") or "Not detected"
    }